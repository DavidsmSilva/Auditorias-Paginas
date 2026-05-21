#!/usr/bin/env python3
"""
docx_report.py — Generador de Reportes DOCX Profesionales
=========================================================
Lee el reporte JSON generado por PaginasAudit y produce un documento
Word (.docx) profesional con portada, resumen ejecutivo, hallazgos
detallados por severidad, fases, recomendaciones y anexos.

Uso:
    python3 docx_report.py <audit-report.json> [output.docx]

Requiere: python-docx (pip install python-docx)
"""

import json
import sys
import os
from datetime import datetime

# ---- Try to import python-docx -------------------------------------------
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


# ===========================================================================
# CONSTANTS — Color Palette
# ===========================================================================
COLOR_PRIMARY = RGBColor(0x0D, 0x11, 0x17)      # Dark background
COLOR_SECONDARY = RGBColor(0x58, 0xA6, 0xFF)     # Blue accent
COLOR_CRITICAL = RGBColor(0xDC, 0x35, 0x45)      # Red
COLOR_HIGH = RGBColor(0xFD, 0x7E, 0x14)          # Orange
COLOR_MEDIUM = RGBColor(0xFF, 0xC1, 0x07)        # Yellow
COLOR_LOW = RGBColor(0x0D, 0xCA, 0xF0)           # Cyan
COLOR_INFO = RGBColor(0x6C, 0x75, 0x7D)          # Gray
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_BG_LIGHT = RGBColor(0xF6, 0xF8, 0xFA)      # Light gray bg

SEVERITY_COLORS = {
    'CRITICAL': COLOR_CRITICAL,
    'HIGH': COLOR_HIGH,
    'MEDIUM': COLOR_MEDIUM,
    'LOW': COLOR_LOW,
    'INFO': COLOR_INFO,
}

SEVERITY_EMOJI = {
    'CRITICAL': '🛑',
    'HIGH': '⚠️ ',
    'MEDIUM': '⚡',
    'LOW': 'ℹ️ ',
    'INFO': '📌',
}


# ===========================================================================
# HELPERS
# ===========================================================================
def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, color=None, size=None, alignment=None):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    # Reduce margins
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_styled_heading(doc, text, level=1):
    """Add a heading with custom color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = COLOR_PRIMARY
        elif level == 2:
            run.font.color.rgb = COLOR_SECONDARY
        elif level == 3:
            run.font.color.rgb = COLOR_PRIMARY
    return h


def add_severity_badge(doc, severity):
    """Add a colored severity badge paragraph."""
    color = SEVERITY_COLORS.get(severity, COLOR_INFO)
    emoji = SEVERITY_EMOJI.get(severity, '❓')
    p = doc.add_paragraph()
    run = p.add_run(f'{emoji}  {severity}')
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(12)
    return p


def fmt_time(seconds):
    """Format seconds to mm:ss or H:MM:SS."""
    seconds = int(seconds or 0)
    h, m = divmod(seconds, 3600)
    m, s = divmod(m, 60)
    if h:
        return f"{h}h {m}m {s}s"
    return f"{m}m {s}s"


# ===========================================================================
# REPORT GENERATOR
# ===========================================================================
class AuditDocxReport:
    """Generates a professional DOCX audit report from JSON data."""

    def __init__(self, json_path, output_path=None):
        self.json_path = json_path
        self.output_path = output_path or os.path.join(
            os.path.dirname(json_path), 'audit-report.docx'
        )
        self.data = self._load_json()
        self.doc = Document()

    def _load_json(self):
        """Load and validate the audit JSON report."""
        with open(self.json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        required = ['meta', 'target', 'summary', 'findings']
        for key in required:
            if key not in data:
                raise ValueError(f"JSON report missing required key: {key}")
        return data

    def _set_document_defaults(self):
        """Set default font and margins."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.color.rgb = COLOR_BLACK

        # Page margins
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    # ---- COVER PAGE ------------------------------------------------------

    def build_cover(self):
        """Build a professional cover page."""
        # Add vertical spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('PaginasAudit')
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = COLOR_SECONDARY

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Reporte de Auditoría de Seguridad')
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_PRIMARY

        # Divider
        self.doc.add_paragraph()
        divider = self.doc.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = divider.add_run('─' * 50)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Target info
        target = self.data.get('target', {})
        meta = self.data.get('meta', {})
        summary = self.data.get('summary', {})

        info_items = [
            ('URL Objetivo', target.get('url', 'N/A')),
            ('Dominio', target.get('domain', 'N/A')),
            ('Dirección IP', target.get('ip', 'N/A')),
            ('Fecha del Reporte', meta.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))),
            ('Estado', target.get('status', 'N/A')),
            ('WAF Detectado', target.get('waf', 'N/A')),
        ]

        info_table = self.doc.add_table(rows=len(info_items), cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (key, val) in enumerate(info_items):
            set_cell_text(info_table.rows[i].cells[0], key, bold=True, size=10)
            set_cell_text(info_table.rows[i].cells[1], val, size=10)
            set_cell_shading(info_table.rows[i].cells[0], 'E8F0FE')

        self.doc.add_paragraph()

        # Risk Score
        risk_score = self._calculate_risk_score()
        risk_label = self._risk_label(risk_score)

        risk_p = self.doc.add_paragraph()
        risk_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = risk_p.add_run(f'Risk Score: {risk_score} — {risk_label}')
        run.bold = True
        run.font.size = Pt(16)
        if risk_score > 50:
            run.font.color.rgb = COLOR_CRITICAL
        elif risk_score > 25:
            run.font.color.rgb = COLOR_HIGH
        elif risk_score > 10:
            run.font.color.rgb = COLOR_MEDIUM
        else:
            run.font.color.rgb = COLOR_LOW

        # Finding counts
        summary_counts = self.doc.add_paragraph()
        summary_counts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        counts_text = '  |  '.join([
            f"🛑 {summary.get('critical', 0)} CRÍTICOS",
            f"⚠️  {summary.get('high', 0)} ALTOS",
            f"⚡ {summary.get('medium', 0)} MEDIOS",
            f"ℹ️  {summary.get('low', 0)} BAJOS",
            f"📌 {summary.get('info', 0)} INFO",
        ])
        run = summary_counts.add_run(counts_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Page break
        self.doc.add_page_break()

    def _calculate_risk_score(self):
        """Calculate weighted risk score."""
        summary = self.data.get('summary', {})
        return (
            summary.get('critical', 0) * 10
            + summary.get('high', 0) * 5
            + summary.get('medium', 0) * 2
            + summary.get('low', 0) * 1
        )

    def _risk_label(self, score):
        if score > 50:
            return 'CRÍTICO'
        elif score > 25:
            return 'ALTO'
        elif score > 10:
            return 'MEDIO'
        return 'BAJO'

    # ---- TABLE OF CONTENTS -----------------------------------------------

    def build_toc(self):
        """Build a simple table of contents."""
        add_styled_heading(self.doc, 'Índice', level=1)

        toc_items = [
            ('1.', 'Resumen Ejecutivo'),
            ('2.', 'Información del Target'),
            ('3.', 'Hallazgos por Severidad'),
            ('', '  3.1  Hallazgos CRÍTICOS'),
            ('', '  3.2  Hallazgos ALTOS'),
            ('', '  3.3  Hallazgos MEDIOS'),
            ('', '  3.4  Hallazgos BAJOS'),
            ('', '  3.5  Hallazgos INFO'),
            ('4.', 'Detalle por Fase de Auditoría'),
            ('', '  4.1  Fase 1 — Assessment'),
            ('', '  4.2  Fase 2 — Malware Analysis'),
            ('', '  4.3  Fase 3 — Brand Protection'),
            ('', '  4.4  Fase 4 — Incident Response'),
            ('5.', 'Recomendaciones'),
            ('6.', 'Anexos'),
        ]

        for num, title in toc_items:
            p = self.doc.add_paragraph()
            run = p.add_run(f'{num} {title}')
            if num and num[-1] == '.':
                run.bold = True
                run.font.size = Pt(11)
            else:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        self.doc.add_page_break()

    # ---- EXECUTIVE SUMMARY -----------------------------------------------

    def build_executive_summary(self):
        """Build the executive summary section."""
        add_styled_heading(self.doc, '1. Resumen Ejecutivo', level=1)

        summary = self.data.get('summary', {})
        total = summary.get('total', 0)
        risk_score = self._calculate_risk_score()
        risk_label = self._risk_label(risk_score)

        p = self.doc.add_paragraph()
        p.add_run(
            f'Se realizó una auditoría de seguridad automatizada contra el target '
            f'{self.data["target"].get("url", "N/A")}. '
            f'Se ejecutaron las 4 fases del pipeline de auditoría, '
            f'identificando un total de {total} hallazgos '
            f'con un Risk Score de {risk_score} ({risk_label}).'
        )

        # Summary table
        self.doc.add_paragraph()
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Light Shading Accent 1'

        severity_data = [
            ('CRÍTICOS', summary.get('critical', 0), 'dc3545'),
            ('ALTOS', summary.get('high', 0), 'fd7e14'),
            ('MEDIOS', summary.get('medium', 0), 'ffc107'),
            ('BAJOS', summary.get('low', 0), '0dcaf0'),
            ('INFO', summary.get('info', 0), '6c757d'),
        ]

        set_cell_text(table.rows[0].cells[0], 'Severidad', bold=True, size=10)
        set_cell_text(table.rows[0].cells[1], 'Cantidad', bold=True, size=10)
        set_cell_shading(table.rows[0].cells[0], '0D1117')
        set_cell_shading(table.rows[0].cells[1], '0D1117')
        table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = COLOR_WHITE
        table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = COLOR_WHITE

        for i, (sev, count, color_hex) in enumerate(severity_data):
            row = table.rows[i + 1]
            sev_color = SEVERITY_COLORS.get(sev, COLOR_BLACK)
            set_cell_text(row.cells[0], f'{SEVERITY_EMOJI.get(sev, "")} {sev}', bold=True, size=10, color=sev_color)
            set_cell_text(row.cells[1], str(count), size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Total row
        total_row = table.rows[5]
        set_cell_text(total_row.cells[0], 'TOTAL', bold=True, size=10)
        set_cell_text(total_row.cells[1], str(total), bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(total_row.cells[0], 'E8F0FE')
        set_cell_shading(total_row.cells[1], 'E8F0FE')

        # Risk score interpretation
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run('Risk Score: ').bold = True
        p.add_run(f'{risk_score} — {risk_label}').bold = True

        interpretations = {
            'CRÍTICO': 'El target presenta vulnerabilidades críticas que requieren atención inmediata. '
                        'Se recomienda una revisión de seguridad completa y la remediación urgente de los hallazgos críticos.',
            'ALTO': 'El target tiene vulnerabilidades significativas que deben ser abordadas prioritariamente. '
                    'Existe riesgo de compromiso si no se toman medidas.',
            'MEDIO': 'El target tiene algunas debilidades de seguridad que deben ser corregidas, '
                     'aunque el riesgo inmediato es moderado.',
            'BAJO': 'El target presenta un perfil de riesgo bajo. Las vulnerabilidades encontradas son '
                    'principalmente informativas o de baja criticidad.',
        }
        p2 = self.doc.add_paragraph()
        p2.add_run(interpretations.get(risk_label, ''))

        self.doc.add_page_break()

    # ---- TARGET INFO -----------------------------------------------------

    def build_target_info(self):
        """Build target information section."""
        add_styled_heading(self.doc, '2. Información del Target', level=1)

        target = self.data.get('target', {})
        timing = self.data.get('timing', {})
        meta = self.data.get('meta', {})

        info_table = self.doc.add_table(rows=8, cols=2)
        info_rows = [
            ('URL', target.get('url', 'N/A')),
            ('Dominio', target.get('domain', 'N/A')),
            ('Dirección IP', target.get('ip', 'N/A')),
            ('Estado', target.get('status', 'N/A')),
            ('WAF', target.get('waf', 'N/A')),
            ('Tecnologías', target.get('technologies', 'N/A')[:120]),
            ('Fecha', meta.get('date', 'N/A')),
            ('Versión', meta.get('version', '1.0.0')),
        ]

        for i, (key, val) in enumerate(info_rows):
            set_cell_text(info_table.rows[i].cells[0], key, bold=True, size=10)
            set_cell_text(info_table.rows[i].cells[1], val, size=10)
            set_cell_shading(info_table.rows[i].cells[0], 'E8F0FE')

        # Timing
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run('Tiempos de Ejecución:').bold = True

        phase_names = {
            'assessment': 'Assessment (Escaneo activo)',
            'malware': 'Malware Analysis',
            'brand': 'Brand Protection',
            'incident': 'Incident Response',
        }
        total_seconds = sum(timing.get(k, 0) for k in phase_names)
        for phase_key, phase_label in phase_names.items():
            secs = timing.get(phase_key, 0)
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(f'{phase_label}: ').bold = True
            p.add_run(fmt_time(secs))

        p = self.doc.add_paragraph(style='List Bullet')
        p.add_run('Total: ').bold = True
        p.add_run(fmt_time(total_seconds))

        self.doc.add_page_break()

    # ---- FINDINGS BY SEVERITY --------------------------------------------

    def build_findings(self):
        """Build findings section organized by severity."""
        add_styled_heading(self.doc, '3. Hallazgos por Severidad', level=1)

        findings = self.data.get('findings', [])
        if not findings:
            self.doc.add_paragraph('No se encontraron hallazgos durante la auditoría.')
            return

        severity_order = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFO']
        severity_count = {
            'CRITICAL': 0, 'HIGH': 0, 'MEDIUM': 0, 'LOW': 0, 'INFO': 0
        }
        severity_findings = {s: [] for s in severity_order}

        for finding in findings:
            sev = finding.get('severity', 'INFO').upper()
            if sev in severity_findings:
                severity_findings[sev].append(finding)
                severity_count[sev] = severity_count.get(sev, 0) + 1

        finding_num = 0
        for sev in severity_order:
            items = severity_findings[sev]
            if not items:
                continue

            # Severity heading
            add_severity_badge(self.doc, sev)

            for finding in items:
                finding_num += 1
                title = finding.get('title', 'Unknown')
                source = finding.get('source', '')
                detail = finding.get('detail', '')
                recommendation = finding.get('recommendation', '')

                # Finding card
                p = self.doc.add_paragraph()
                run = p.add_run(f'#{finding_num}  {title}')
                run.bold = True
                run.font.size = Pt(11)
                if sev in SEVERITY_COLORS:
                    run.font.color.rgb = SEVERITY_COLORS[sev]

                # Source tag
                if source:
                    p2 = self.doc.add_paragraph()
                    run2 = p2.add_run(f'Fuente: {source}')
                    run2.font.size = Pt(9)
                    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    run2.italic = True
                    p2.paragraph_format.space_before = Pt(0)
                    p2.paragraph_format.space_after = Pt(2)

                # Detail
                if detail:
                    p3 = self.doc.add_paragraph(detail)
                    p3.paragraph_format.space_before = Pt(4)
                    p3.paragraph_format.space_after = Pt(4)

                # Recommendation
                if recommendation:
                    p4 = self.doc.add_paragraph()
                    run4 = p4.add_run(f'Recomendación: {recommendation}')
                    run4.font.size = Pt(9)
                    run4.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

                # Evidence (if present)
                evidence = finding.get('evidence', '')
                if evidence:
                    p5 = self.doc.add_paragraph()
                    run5 = p5.add_run(f'Evidencia: {evidence[:200]}')
                    run5.font.size = Pt(8)
                    run5.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    run5.italic = True

                # Thin separator
                self.doc.add_paragraph()

            self.doc.add_paragraph()

        self.doc.add_page_break()

    # ---- PHASE DETAILS ---------------------------------------------------

    def build_phase_details(self):
        """Build detailed phase analysis section."""
        add_styled_heading(self.doc, '4. Detalle por Fase de Auditoría', level=1)

        phases = [
            ('4.1', 'Assessment — Escaneo Activo',
             'Esta fase incluyó escaneo de puertos (Nmap), fingerprinting web (WhatWeb), '
             'escaneo de vulnerabilidades (Nikto, Nuclei), evaluación SSL/TLS (sslscan, testssl.sh), '
             'enumeración DNS (dnsrecon) y fuzzing de directorios (Gobuster).',
             'scans'),
            ('4.2', 'Malware Analysis — Análisis de Código y Dependencias',
             'Esta fase incluyó análisis de cabeceras de seguridad HTTP, detección de metadatos '
             'en recursos descargados (ExifTool), escaneo de patrones de malware (YARA), y '
             'revisión de configuración de seguridad del servidor web.',
             'malware'),
            ('4.3', 'Brand Protection — Protección de Marca',
             'Esta fase incluyó detección de typosquatting (dnstwist), recolección OSINT '
             '(theHarvester), enumeración de subdominios (Sublist3r), y verificación de '
             'fugas de credenciales en brechas públicas (HIBP).',
             'osint'),
            ('4.4', 'Incident Response — Preparación ante Incidentes',
             'Esta fase evaluó la preparación técnica para responder a incidentes de seguridad, '
             'incluyendo disponibilidad de herramientas forenses, captura de tráfico, '
             'monitoreo de integridad y capacidad de recuperación.',
             'ir'),
        ]

        for num, title, desc, _ in phases:
            add_styled_heading(self.doc, f'{num}  {title}', level=2)
            self.doc.add_paragraph(desc)

            # Get related findings
            source_map = {
                '4.1': ['Nmap', 'Nikto', 'WhatWeb', 'Nuclei', 'SSLScan', 'testssl',
                         'DNSRecon', 'DNS', 'Gobuster', 'WPScan', 'Security Headers'],
                '4.2': ['ExifTool', 'YARA', 'Security Headers', 'Pre-flight'],
                '4.3': ['dnstwist', 'theHarvester', 'Sublist3r', 'HIBP'],
                '4.4': ['IR Readiness'],
            }

            related = source_map.get(num, [])
            phase_findings = [
                f for f in self.data.get('findings', [])
                if any(rs.lower() in f.get('source', '').lower() for rs in related)
            ]

            if phase_findings:
                p = self.doc.add_paragraph()
                p.add_run(f'Hallazgos relacionados: {len(phase_findings)}').bold = True
                for finding in phase_findings[:5]:  # Show max 5
                    sev = finding.get('severity', 'INFO')
                    color = SEVERITY_COLORS.get(sev, COLOR_BLACK)
                    title = finding.get('title', '')
                    p2 = self.doc.add_paragraph(style='List Bullet')
                    run = p2.add_run(f'[{sev}] ')
                    run.bold = True
                    run.font.color.rgb = color
                    p2.add_run(title)
            else:
                self.doc.add_paragraph('Sin hallazgos específicos en esta fase.')

            self.doc.add_paragraph()

    # ---- RECOMMENDATIONS -------------------------------------------------

    def build_recommendations(self):
        """Build recommendations section."""
        add_styled_heading(self.doc, '5. Recomendaciones', level=1)

        # Collect unique recommendations from findings
        findings = self.data.get('findings', [])
        severity_order = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFO']

        recommendations = {s: [] for s in severity_order}

        for finding in findings:
            sev = finding.get('severity', 'INFO').upper()
            rec = finding.get('recommendation', '').strip()
            if rec:
                recommendations[sev].append(rec)

        # Grouped recommendations
        for sev in severity_order:
            recs = recommendations[sev]
            if not recs:
                continue

            add_severity_badge(self.doc, sev)

            # Unique recommendations
            seen = set()
            for rec in recs:
                if rec not in seen:
                    seen.add(rec)
                    p = self.doc.add_paragraph(style='List Bullet')
                    p.add_run(rec)

            self.doc.add_paragraph()

        # General recommendations
        self.doc.add_paragraph()
        add_styled_heading(self.doc, 'Recomendaciones Generales', level=2)
        general_recs = [
            'Mantener todo el software actualizado, incluyendo CMS, plugins, librerías y sistema operativo.',
            'Implementar un programa de bug bounty o pruebas de penetración periódicas.',
            'Configurar monitoreo de seguridad continuo con alertas en tiempo real.',
            'Establecer un proceso formal de respuesta a incidentes.',
            'Realizar backups periódicos y probar la restauración al menos trimestralmente.',
            'Implementar autenticación multifactor (MFA) en todos los accesos administrativos.',
            'Mantener un inventario actualizado de activos digitales y superficies de ataque.',
        ]

        for rec in general_recs:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(rec)

        self.doc.add_page_break()

    # ---- APPENDIX --------------------------------------------------------

    def build_appendix(self):
        """Build appendix with technical details."""
        add_styled_heading(self.doc, '6. Anexos', level=1)

        # Tools used
        add_styled_heading(self.doc, '6.1  Herramientas Utilizadas', level=2)
        tools_by_phase = {
            'Assessment': ['nmap', 'whatweb', 'nikto', 'sslscan', 'testssl.sh',
                           'dnsrecon', 'nuclei', 'gobuster', 'wpscan'],
            'Malware Analysis': ['curl', 'exiftool', 'yara'],
            'Brand Protection': ['dnstwist', 'theharvester', 'sublist3r', 'hibp-check'],
            'Incident Response': ['tcpdump', 'tshark', 'volatility', 'sleuthkit',
                                  'aide', 'rsync'],
        }

        for phase, tools in tools_by_phase.items():
            p = self.doc.add_paragraph()
            p.add_run(f'{phase}: ').bold = True
            p.add_run(', '.join(tools))

        # Timing details
        self.doc.add_paragraph()
        add_styled_heading(self.doc, '6.2  Tiempos de Ejecución', level=2)
        timing = self.data.get('timing', {})
        phases = ['assessment', 'malware', 'brand', 'incident']
        phase_labels = ['Assessment', 'Malware Analysis', 'Brand Protection', 'Incident Response']

        time_table = self.doc.add_table(rows=len(phases) + 2, cols=2)
        set_cell_text(time_table.rows[0].cells[0], 'Fase', bold=True, size=10)
        set_cell_text(time_table.rows[0].cells[1], 'Duración', bold=True, size=10)
        set_cell_shading(time_table.rows[0].cells[0], '0D1117')
        set_cell_shading(time_table.rows[0].cells[1], '0D1117')
        time_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = COLOR_WHITE
        time_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = COLOR_WHITE

        total_secs = 0
        for i, (pk, pl) in enumerate(zip(phases, phase_labels)):
            secs = timing.get(pk, 0)
            total_secs += secs
            set_cell_text(time_table.rows[i + 1].cells[0], pl, size=10)
            set_cell_text(time_table.rows[i + 1].cells[1], fmt_time(secs), size=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

        set_cell_text(time_table.rows[-1].cells[0], 'TOTAL', bold=True, size=10)
        set_cell_text(time_table.rows[-1].cells[1], fmt_time(total_secs), bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(time_table.rows[-1].cells[0], 'E8F0FE')
        set_cell_shading(time_table.rows[-1].cells[1], 'E8F0FE')

        # Footer
        self.doc.add_paragraph()
        footer_p = self.doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(
            '— Generado por PaginasAudit —\n'
            'https://github.com/DavidsmSilva/Auditorias-Paginas'
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

    # ---- MAIN BUILD ------------------------------------------------------

    def build(self):
        """Build the complete DOCX report."""
        print("[DOCX] Generando reporte profesional...")

        self._set_document_defaults()
        self.build_cover()
        self.build_toc()
        self.build_executive_summary()
        self.build_target_info()
        self.build_findings()
        self.build_phase_details()
        self.build_recommendations()
        self.build_appendix()

        # Save
        self.doc.save(self.output_path)
        print(f"[DOCX] Reporte generado: {self.output_path}")
        return self.output_path


# ===========================================================================
# CLI ENTRY POINT
# ===========================================================================
def main():
    if len(sys.argv) < 2:
        print("Uso: python3 docx_report.py <audit-report.json> [output.docx]")
        print("")
        print("Genera un reporte DOCX profesional desde el JSON de auditoría.")
        sys.exit(1)

    json_path = sys.argv[1]
    if not os.path.isfile(json_path):
        print(f"Error: No se encuentra el archivo JSON: {json_path}")
        sys.exit(1)

    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not DOCX_AVAILABLE:
        print("[DOCX] ERROR: python-docx no está instalado.")
        print("[DOCX] Instale con: pip3 install python-docx")
        print("[DOCX] O: sudo apt install python3-docx")
        sys.exit(1)

    try:
        report = AuditDocxReport(json_path, output_path)
        output = report.build()
        print(f"[DOCX] ✅ Reporte DOCX generado exitosamente: {output}")
    except Exception as e:
        print(f"[DOCX] ❌ Error generando reporte: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
